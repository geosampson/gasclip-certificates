#!/usr/bin/env python3
"""
GasClip Certificate Generator v5.0 - Word-based Edition
Generates gas detector calibration certificates using Word templates
"""

import tkinter as tk
from tkinter import ttk, messagebox, filedialog
from datetime import datetime, timedelta
from pathlib import Path
import subprocess
import shutil
from docx import Document

class GasClipCertificateGenerator:
    def __init__(self, root):
        self.root = root
        self.root.title("GasClip Certificate Generator v5.0 (Word Edition)")
        self.root.geometry("900x750")
        
        # Product information
        self.products = {
            "D4PQ": {
                "name": "MGC-S+ (MGC-SIMPLEPLUS)",
                "template": "D4PQ236599_clean.docx",
                "detector_life_months": 36,
                "calibration_days": 1095
            },
            "SOSP": {
                "name": "SGC-O (Single Gas Clip O2)",
                "template": "SOSP215459_clean.docx",
                "detector_life_months": 24,
                "calibration_days": 730
            },
            "SCSQ": {
                "name": "SGC-C (Single Gas Clip CO)",
                "template": "SCSQ175392_clean.docx",
                "detector_life_months": 24,
                "calibration_days": 730
            },
            "D4SQ": {
                "name": "MGC-S (MGC-SIMPLE)",
                "template": "D4SQ106733_clean.docx",
                "detector_life_months": 36,
                "calibration_days": 1095
            },
            "SHSP": {
                "name": "SGC-H (Single Gas Clip H2S)",
                "template": "SHSP085112_clean.docx",
                "detector_life_months": 24,
                "calibration_days": 730
            }
        }
        
        # Certificate counter
        self.certificates_generated = 0
        self.generated_files = []
        
        # Create output directory
        self.output_dir = Path("output")
        self.output_dir.mkdir(exist_ok=True)
        
        self.create_widgets()
        
    def create_widgets(self):
        # Title
        title_label = tk.Label(
            self.root,
            text="GasClip Certificate Generator v5.0",
            font=("Arial", 18, "bold"),
            fg="#1e3a8a"
        )
        title_label.pack(pady=15)
        
        # Main frame
        main_frame = ttk.Frame(self.root, padding="20")
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        # Product selection
        ttk.Label(main_frame, text="Select Product:", font=("Arial", 11, "bold")).grid(
            row=0, column=0, sticky=tk.W, pady=5
        )
        
        self.product_var = tk.StringVar()
        self.product_combo = ttk.Combobox(
            main_frame,
            textvariable=self.product_var,
            values=[f"{k} - {v['name']}" for k, v in self.products.items()],
            state="readonly",
            width=50,
            font=("Arial", 10)
        )
        self.product_combo.grid(row=0, column=1, pady=5, sticky=tk.EW)
        self.product_combo.bind("<<ComboboxSelected>>", self.on_product_selected)
        
        # Serial Number Prefix (display only)
        ttk.Label(main_frame, text="Serial Number Prefix:", font=("Arial", 11)).grid(
            row=1, column=0, sticky=tk.W, pady=5
        )
        self.prefix_label = tk.Label(
            main_frame,
            text="",
            font=("Arial", 14, "bold"),
            fg="#2563eb"
        )
        self.prefix_label.grid(row=1, column=1, sticky=tk.W, pady=5)
        
        # Serial Number (digits only)
        ttk.Label(main_frame, text="Serial Number (digits only):", font=("Arial", 11)).grid(
            row=2, column=0, sticky=tk.W, pady=5
        )
        self.serial_entry = ttk.Entry(main_frame, width=52, font=("Arial", 10))
        self.serial_entry.grid(row=2, column=1, pady=5, sticky=tk.EW)
        self.serial_entry.bind("<KeyRelease>", self.validate_serial)
        self.serial_entry.bind("<Return>", lambda e: self.activation_entry.focus())
        
        ttk.Label(main_frame, text="Example: 175392", font=("Arial", 9), foreground="gray").grid(
            row=3, column=1, sticky=tk.W
        )
        
        # Activation Date
        ttk.Label(main_frame, text="Activation Date:", font=("Arial", 11)).grid(
            row=4, column=0, sticky=tk.W, pady=5
        )
        self.activation_entry = ttk.Entry(main_frame, width=52, font=("Arial", 10))
        self.activation_entry.grid(row=4, column=1, pady=5, sticky=tk.EW)
        self.activation_entry.bind("<KeyRelease>", self.format_date)
        self.activation_entry.bind("<Return>", lambda e: self.lot_entry.focus())
        
        ttk.Label(main_frame, text="Type: 26022025 → Auto-formats to: 26/02/2025", font=("Arial", 9), foreground="gray").grid(
            row=5, column=1, sticky=tk.W
        )
        
        # Lot Number
        ttk.Label(main_frame, text="Lot Number:", font=("Arial", 11)).grid(
            row=6, column=0, sticky=tk.W, pady=5
        )
        self.lot_entry = ttk.Entry(main_frame, width=52, font=("Arial", 10))
        self.lot_entry.grid(row=6, column=1, pady=5, sticky=tk.EW)
        self.lot_entry.bind("<Return>", lambda e: self.gas_prod_entry.focus())
        
        ttk.Label(main_frame, text="Example: RR2310181807 or 25-3348", font=("Arial", 9), foreground="gray").grid(
            row=7, column=1, sticky=tk.W
        )
        
        # Gas Production Date
        ttk.Label(main_frame, text="Gas Production Date:", font=("Arial", 11)).grid(
            row=8, column=0, sticky=tk.W, pady=5
        )
        self.gas_prod_entry = ttk.Entry(main_frame, width=52, font=("Arial", 10))
        self.gas_prod_entry.grid(row=8, column=1, pady=5, sticky=tk.EW)
        self.gas_prod_entry.bind("<KeyRelease>", self.format_date)
        self.gas_prod_entry.bind("<Return>", lambda e: self.calibration_entry.focus())
        
        ttk.Label(main_frame, text="Type: 19102023", font=("Arial", 9), foreground="gray").grid(
            row=9, column=1, sticky=tk.W
        )
        
        # Calibration Date
        ttk.Label(main_frame, text="Calibration Date:", font=("Arial", 11)).grid(
            row=10, column=0, sticky=tk.W, pady=5
        )
        self.calibration_entry = ttk.Entry(main_frame, width=52, font=("Arial", 10))
        self.calibration_entry.grid(row=10, column=1, pady=5, sticky=tk.EW)
        self.calibration_entry.bind("<KeyRelease>", self.format_date)
        self.calibration_entry.bind("<Return>", lambda e: self.generate_certificate())
        
        ttk.Label(main_frame, text="Type: 26022024", font=("Arial", 9), foreground="gray").grid(
            row=11, column=1, sticky=tk.W
        )
        
        # Buttons frame
        button_frame = ttk.Frame(main_frame)
        button_frame.grid(row=12, column=0, columnspan=2, pady=20)
        
        # Generate button
        self.generate_btn = ttk.Button(
            button_frame,
            text="✓ Generate Certificate",
            command=self.generate_certificate,
            width=25
        )
        self.generate_btn.pack(side=tk.LEFT, padx=5)
        
        # Clear button
        clear_btn = ttk.Button(
            button_frame,
            text="Clear Form",
            command=self.clear_form,
            width=15
        )
        clear_btn.pack(side=tk.LEFT, padx=5)
        
        # Status label
        self.status_label = tk.Label(
            main_frame,
            text="Ready to generate certificates",
            font=("Arial", 10),
            fg="#059669"
        )
        self.status_label.grid(row=13, column=0, columnspan=2, pady=10)
        
        # Counter
        self.counter_label = tk.Label(
            main_frame,
            text="Certificates Generated: 0",
            font=("Arial", 12, "bold"),
            fg="#16a34a"
        )
        self.counter_label.grid(row=14, column=0, columnspan=2, pady=5)
        
        # Finish button
        self.finish_btn = ttk.Button(
            main_frame,
            text="📁 Finish All Forms & Create Invoice Folder",
            command=self.finish_all_forms,
            width=40
        )
        self.finish_btn.grid(row=15, column=0, columnspan=2, pady=10)
        
        # Tip label
        tip_label = tk.Label(
            main_frame,
            text="💡 Tip: Press Enter to move to next field, ↑↓ arrows to navigate",
            font=("Arial", 9),
            fg="#6b7280"
        )
        tip_label.grid(row=16, column=0, columnspan=2, pady=5)
        
        # Configure grid weights
        main_frame.columnconfigure(1, weight=1)
        
    def on_product_selected(self, event=None):
        """Update prefix when product is selected"""
        selection = self.product_var.get()
        if selection:
            prefix = selection.split(" - ")[0]
            self.prefix_label.config(text=f"{prefix}XXXXXX")
            
    def validate_serial(self, event=None):
        """Validate serial number input (digits only)"""
        value = self.serial_entry.get()
        # Remove non-digits
        cleaned = ''.join(c for c in value if c.isdigit())
        if cleaned != value:
            self.serial_entry.delete(0, tk.END)
            self.serial_entry.insert(0, cleaned)
            
    def format_date(self, event=None):
        """Auto-format date as user types"""
        widget = event.widget
        value = widget.get().replace("/", "")  # Remove existing slashes
        
        # Only keep digits
        value = ''.join(c for c in value if c.isdigit())
        
        # Format as DD/MM/YYYY
        if len(value) >= 2:
            formatted = value[:2]
            if len(value) >= 4:
                formatted += "/" + value[2:4]
                if len(value) >= 8:
                    formatted += "/" + value[4:8]
                elif len(value) > 4:
                    formatted += "/" + value[4:]
            elif len(value) > 2:
                formatted += "/" + value[2:]
            
            # Update widget
            widget.delete(0, tk.END)
            widget.insert(0, formatted)
            
    def validate_date(self, date_str):
        """Validate if date is real"""
        try:
            datetime.strptime(date_str, "%d/%m/%Y")
            return True
        except ValueError:
            return False
            
    def generate_certificate(self):
        """Generate certificate using Word template"""
        # Validate inputs
        if not self.product_var.get():
            messagebox.showerror("Error", "Please select a product")
            return
            
        prefix = self.product_var.get().split(" - ")[0]
        serial_digits = self.serial_entry.get().strip()
        
        if not serial_digits:
            messagebox.showerror("Error", "Please enter serial number digits")
            return
            
        if len(serial_digits) < 5:
            messagebox.showerror("Error", "Serial number should be at least 5 digits")
            return
            
        activation = self.activation_entry.get().strip()
        lot = self.lot_entry.get().strip()
        gas_prod = self.gas_prod_entry.get().strip()
        calibration = self.calibration_entry.get().strip()
        
        if not all([activation, lot, gas_prod, calibration]):
            messagebox.showerror("Error", "Please fill in all fields")
            return
            
        # Validate dates
        for date_field, date_value in [("Activation", activation), ("Gas Production", gas_prod), ("Calibration", calibration)]:
            if not self.validate_date(date_value):
                messagebox.showerror("Error", f"Invalid {date_field} date: {date_value}")
                return
                
        # Generate certificate
        try:
            full_serial = f"{prefix}{serial_digits}"
            product_info = self.products[prefix]
            
            self.status_label.config(text=f"Generating certificate for {full_serial}...", fg="#f59e0b")
            self.root.update()
            
            # Process Word template
            template_path = Path(__file__).parent / "templates_word" / product_info["template"]
            
            if not template_path.exists():
                messagebox.showerror("Error", f"Template not found: {template_path}")
                return
                
            # Open Word document
            doc = Document(template_path)
            
            # Replace text in the document
            replacements = self.get_replacements(prefix, full_serial, activation, lot, gas_prod, calibration, product_info)
            
            # Replace in paragraphs
            for paragraph in doc.paragraphs:
                for old_text, new_text in replacements.items():
                    if old_text in paragraph.text:
                        for run in paragraph.runs:
                            if old_text in run.text:
                                run.text = run.text.replace(old_text, new_text)
                                
            # Replace in tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for old_text, new_text in replacements.items():
                                if old_text in paragraph.text:
                                    for run in paragraph.runs:
                                        if old_text in run.text:
                                            run.text = run.text.replace(old_text, new_text)
            
            # Save modified Word document
            output_docx = self.output_dir / f"{full_serial}.docx"
            doc.save(output_docx)
            
            # Convert to PDF using LibreOffice
            output_pdf = self.output_dir / f"{full_serial}.pdf"
            self.convert_to_pdf(output_docx, output_pdf)
            
            # Delete temporary Word file
            output_docx.unlink()
            
            self.certificates_generated += 1
            self.generated_files.append(output_pdf)
            self.counter_label.config(text=f"Certificates Generated: {self.certificates_generated}")
            self.status_label.config(text=f"✓ Certificate generated: {full_serial}.pdf", fg="#16a34a")
            
            # Clear form for next certificate
            self.clear_form(keep_product=True)
            
        except Exception as e:
            messagebox.showerror("Error", f"Failed to generate certificate: {str(e)}")
            self.status_label.config(text="Error generating certificate", fg="#dc2626")
            
    def get_replacements(self, prefix, full_serial, activation, lot, gas_prod, calibration, product_info):
        """Get text replacements based on product template"""
        replacements = {}
        
        # Common replacements
        replacements[full_serial.replace(prefix, prefix)] = full_serial  # This handles the serial number
        
        # Get the template serial number from the template filename
        template_serial = product_info["template"].replace("_clean.docx", "")
        
        # Replace template serial with new serial
        replacements[template_serial] = full_serial
        
        # Get template dates from the document (we'll need to identify these)
        # For now, we'll use a simple approach: replace common date patterns
        
        # Read the template to find dates to replace
        template_path = Path(__file__).parent / "templates_word" / product_info["template"]
        doc = Document(template_path)
        
        # Extract dates from template
        template_dates = []
        for paragraph in doc.paragraphs:
            text = paragraph.text
            # Find dates in format DD/MM/YYYY
            import re
            dates = re.findall(r'\d{2}/\d{2}/\d{4}', text)
            template_dates.extend(dates)
            
        # Map template dates to new dates (in order of appearance)
        # This is a simplified approach - you may need to adjust based on actual template structure
        if len(template_dates) >= 3:
            replacements[template_dates[0]] = activation  # First date is activation
            replacements[template_dates[1]] = gas_prod     # Second is gas production
            replacements[template_dates[2]] = calibration  # Third is calibration
            
        # Replace lot number (find the lot pattern in template)
        for paragraph in doc.paragraphs:
            text = paragraph.text
            # Look for lot number patterns
            if "Lot Number" in text or "25-3348" in text or "RR" in text:
                # Extract the lot number from template
                import re
                lot_match = re.search(r'(25-\d+|RR\d+|[A-Z]{1,3}\s*\d+\s*ppm|O2\s*\d+\s*%|CO\s*\d+\s*ppm|H2S\s*\d+\s*ppm)', text)
                if lot_match:
                    template_lot = lot_match.group(1)
                    replacements[template_lot] = lot
                    break
                    
        return replacements
        
    def convert_to_pdf(self, docx_path, pdf_path):
        """Convert Word document to PDF using LibreOffice"""
        try:
            # Use LibreOffice to convert
            subprocess.run([
                "libreoffice",
                "--headless",
                "--convert-to", "pdf",
                "--outdir", str(pdf_path.parent),
                str(docx_path)
            ], check=True, capture_output=True)
            
            return True
        except Exception as e:
            raise Exception(f"PDF conversion failed: {str(e)}")
            
    def clear_form(self, keep_product=False):
        """Clear all input fields"""
        if not keep_product:
            self.product_var.set("")
            self.prefix_label.config(text="")
        self.serial_entry.delete(0, tk.END)
        self.activation_entry.delete(0, tk.END)
        self.lot_entry.delete(0, tk.END)
        self.gas_prod_entry.delete(0, tk.END)
        self.calibration_entry.delete(0, tk.END)
        self.serial_entry.focus()
        
    def finish_all_forms(self):
        """Organize generated certificates into invoice folder"""
        if self.certificates_generated == 0:
            messagebox.showinfo("Info", "No certificates generated yet")
            return
            
        invoice_number = tk.simpledialog.askstring(
            "Invoice Number",
            f"Enter invoice number for {self.certificates_generated} certificate(s):"
        )
        
        if not invoice_number:
            return
            
        # Create invoice folder
        invoice_folder = self.output_dir / f"Invoice_{invoice_number}"
        invoice_folder.mkdir(exist_ok=True)
        
        # Move all generated files to invoice folder
        for pdf_file in self.generated_files:
            if pdf_file.exists():
                shutil.move(str(pdf_file), str(invoice_folder / pdf_file.name))
                
        messagebox.showinfo(
            "Success",
            f"✓ {self.certificates_generated} certificate(s) organized in:\n{invoice_folder}"
        )
        
        # Reset counter
        self.certificates_generated = 0
        self.generated_files = []
        self.counter_label.config(text="Certificates Generated: 0")
        self.status_label.config(text="Ready to generate certificates", fg="#059669")

def main():
    root = tk.Tk()
    app = GasClipCertificateGenerator(root)
    root.mainloop()

if __name__ == "__main__":
    main()
